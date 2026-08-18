"""Markdown → docx 简易转换器（word-export change）。

支持：标题 1-4 级、段落、有序/无序列表、管道表格、**加粗**、引用块。
其余语法按纯段落处理（见 brief Non-goals）。
"""
from __future__ import annotations

import io
import re

from docx import Document
from fastapi import Response

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_HEADING = re.compile(r"^(#{1,4})\s+(.*)")
_OLIST = re.compile(r"^\d+[.、]\s*(.*)")
_ULIST = re.compile(r"^[-*]\s+(.*)")
_TABLE_SEP = re.compile(r"^\|[\s:\-|]+\|?\s*$")


def _add_runs(paragraph, text: str) -> None:
    """把 **加粗** 语法拆成 bold Run，其余为普通 Run。"""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx_bytes(md: str, title: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue

        # 管道表格：当前行 | 开头 + 下一行是分隔行
        if s.startswith("|") and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1].strip()):
            header = [c.strip() for c in s.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Table Grid"
            for j, h in enumerate(header):
                _add_runs(table.rows[0].cells[j].paragraphs[0], h)
            for ri, row in enumerate(rows, 1):
                for j, cell in enumerate(row[: len(header)]):
                    _add_runs(table.rows[ri].cells[j].paragraphs[0], cell)
            continue

        m = _HEADING.match(s)
        if m:
            _add_runs(doc.add_heading("", level=min(len(m.group(1)), 4)), m.group(2))
            i += 1
            continue

        m = _OLIST.match(s)
        if m:
            _add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        m = _ULIST.match(s)
        if m:
            _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue

        if s.startswith(">"):
            _add_runs(doc.add_paragraph(style="Intense Quote"), s.lstrip("> "))
            i += 1
            continue

        _add_runs(doc.add_paragraph(), s)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_response(data: bytes, ascii_name: str) -> Response:
    return Response(
        content=data,
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{ascii_name}"'},
    )
